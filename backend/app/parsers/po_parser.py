from __future__ import annotations
import re
import logging
from datetime import date
from pathlib import Path
from typing import Optional

from app.schemas import POParseSchema, LineItemSchema

logger = logging.getLogger(__name__)


def _extract_text_pdf(path: str) -> str:
    import pdfplumber
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            parts.append(text)
            if len(text.strip()) < 50:
                try:
                    import pytesseract
                    img = page.to_image(resolution=200).original
                    ocr = pytesseract.image_to_string(img)
                    if len(ocr.strip()) > len(text.strip()):
                        parts[-1] = ocr
                except Exception:
                    pass
    return "\n".join(parts)


def _extract_text_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _first_match(pattern: str, text: str, flags: int = re.IGNORECASE) -> Optional[str]:
    m = re.search(pattern, text, flags)
    return m.group(1).strip() if m else None


def _parse_date(raw: str) -> Optional[date]:
    if not raw:
        return None
    from datetime import datetime
    for fmt in ("%d/%m/%Y", "%d-%m-%Y", "%Y-%m-%d", "%B %d, %Y", "%b %d, %Y", "%d %B %Y"):
        try:
            return datetime.strptime(raw.strip(), fmt).date()
        except ValueError:
            continue
    return None


def _parse_amount(raw: str) -> Optional[float]:
    if not raw:
        return None
    cleaned = re.sub(r"[^\d.]", "", raw)
    try:
        return float(cleaned)
    except ValueError:
        return None


def parse_po(file_path: str) -> POParseSchema:
    ext = Path(file_path).suffix.lower()
    try:
        if ext == ".pdf":
            text = _extract_text_pdf(file_path)
        elif ext in (".docx", ".doc"):
            text = _extract_text_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    except Exception as e:
        logger.error("PO text extraction failed: %s", e)
        return POParseSchema()

    po_number = (
        _first_match(r"P\.?O\.?\s*(?:Number|No\.?|#)\s*[:\-]?\s*([A-Z0-9\-/]+)", text)
        or _first_match(r"Purchase\s*Order\s*(?:No\.?|Number|#)?\s*[:\-]?\s*([A-Z0-9\-/]+)", text)
    )
    customer_name = (
        _first_match(r"(?:Vendor|Supplier|To|Bill\s*To)\s*[:\-]?\s*([A-Za-z][\w\s&.,'-]{2,60})", text)
        or _first_match(r"(?:Company|Firm)\s*[:\-]?\s*([A-Za-z][\w\s&.,'-]{2,60})", text)
    )
    po_date_raw = _first_match(r"(?:PO\s*Date|Order\s*Date|Date\s*of\s*Order)\s*[:\-]?\s*([\d\w/\-, ]+)", text)
    expiry_date_raw = (
        _first_match(r"(?:Expiry|Expiration|Delivery|Completion)\s*Date\s*[:\-]?\s*([\d\w/\-, ]+)", text)
        or _first_match(r"(?:Due|Required)\s*(?:By|Date)\s*[:\-]?\s*([\d\w/\-, ]+)", text)
    )
    total_raw = (
        _first_match(r"(?:Total|Grand\s*Total|PO\s*Value|Contract\s*Value)\s*[:\-]?\s*(?:AED|USD|SAR|INR)?\s*([\d,]+\.?\d*)", text)
        or _first_match(r"(?:Amount|Value)\s*[:\-]?\s*(?:AED|USD|SAR|INR)?\s*([\d,]+\.?\d*)", text)
    )
    ship_to = _first_match(r"Ship\s*To\s*[:\-]?\s*((?:.+\n?){1,4})", text)
    bill_to = _first_match(r"Bill\s*To\s*[:\-]?\s*((?:.+\n?){1,4})", text)
    billing_terms = _first_match(r"(?:Billing|Payment)\s*Terms?\s*[:\-]?\s*([^\n]{3,60})", text)
    payment_terms = _first_match(r"(?:Net\s*\d+|Due\s*on\s*Receipt|[0-9]+\s*days?)[^\n]{0,40}", text)
    signatory = _first_match(r"(?:Authoris[zes]d?\s*(?:By|Signatory)|Approved\s*By)\s*[:\-]?\s*([A-Za-z][\w\s.]{2,50})", text)

    line_items: list[LineItemSchema] = []
    line_pattern = re.findall(
        r"([A-Za-z][\w\s&,.-]{5,80})\s+(\d+(?:\.\d+)?)\s+([\d,]+\.?\d*)\s+([\d,]+\.?\d*)",
        text
    )
    for match in line_pattern[:20]:
        try:
            line_items.append(LineItemSchema(
                service=match[0].strip(),
                qty=float(match[1]),
                rate=_parse_amount(match[2]) or 0,
                amount=_parse_amount(match[3]) or 0,
            ))
        except Exception:
            continue

    return POParseSchema(
        po_number=po_number,
        customer_name=customer_name,
        po_date=_parse_date(po_date_raw) if po_date_raw else None,
        expiry_date=_parse_date(expiry_date_raw) if expiry_date_raw else None,
        total_value=_parse_amount(total_raw) if total_raw else None,
        billing_terms=billing_terms,
        payment_terms=payment_terms,
        ship_to_address=ship_to,
        bill_to_address=bill_to,
        authorised_signatory=signatory,
        line_items=line_items,
    )
