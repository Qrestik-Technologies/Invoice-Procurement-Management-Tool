from __future__ import annotations
import re
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

from app.schemas import InvoiceParseSchema, LineItemSchema

logger = logging.getLogger(__name__)

REQUIRED_FIELDS = ["invoice_number", "invoice_date", "total", "customer_name"]

# ---------------------------------------------------------------------------
# Vendor signatures — order matters; more-specific first
# ---------------------------------------------------------------------------
VENDOR_SIGNATURES: dict[str, list[str]] = {
    "qrestik":  ["qrestik", "hor al anz", "nrakaeak", "rakbank"],
    "infinitum": ["infinitum global", "infinitumglobal.org", "jp morgan chase"],
}

# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------

def _extract_text_pdf(path: str) -> str:
    import pdfplumber
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            parts.append(text)
            # OCR fallback for image-based pages
            if len(text.strip()) < 50:
                try:
                    import pytesseract
                    img = page.to_image(resolution=200).original
                    ocr = pytesseract.image_to_string(img)
                    if len(ocr.strip()) > len(text.strip()):
                        parts[-1] = ocr
                except Exception:
                    pass
    return "\n".join(parts)


def _extract_text_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Vendor detection
# ---------------------------------------------------------------------------

def _detect_vendor(text: str) -> str:
    t = text.lower()
    for vendor, patterns in VENDOR_SIGNATURES.items():
        if any(p in t for p in patterns):
            return vendor
    return "generic"


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

def _first_match(pattern: str, text: str, flags: int = re.IGNORECASE) -> Optional[str]:
    m = re.search(pattern, text, flags)
    return m.group(1).strip() if m else None


def _parse_date(raw: str) -> Optional[object]:
    """Try multiple date formats, return date or None."""
    if not raw:
        return None
    # Strip ordinal suffixes: 1st → 1
    cleaned = re.sub(r"(\d+)(?:st|nd|rd|th)", r"\1", raw).strip()
    for fmt in (
        "%d/%m/%Y", "%m/%d/%Y", "%d-%m-%Y", "%m-%d-%Y",
        "%B %d, %Y", "%b %d, %Y", "%d %B %Y", "%d %b %Y",
    ):
        try:
            return datetime.strptime(cleaned, fmt).date()
        except ValueError:
            continue
    return None


# ---------------------------------------------------------------------------
# Qrestik-specific extraction
# ---------------------------------------------------------------------------

def _extract_qrestik(text: str) -> dict:
    data: dict = {}

    # Invoice number — "Invoice number : 0069"
    m = re.search(r"Invoice\s+number\s*[:\-]\s*(\d+)", text, re.I)
    if m:
        data["invoice_number"] = m.group(1).strip()

    # Date — "Date : 13/03/2026"
    m = re.search(r"Date\s*[:\-]\s*(\d{1,2}/\d{1,2}/\d{4})", text, re.I)
    if m:
        data["invoice_date"] = _parse_date(m.group(1))

    data["currency"] = "AED"
    data["vendor_name"] = "Qrestik Technologies L.L.C"

    # Total — "AED 7,345"
    m = re.search(r"AED\s+([\d,]+(?:\.\d+)?)$", text, re.MULTILINE)
    if not m:
        m = re.search(r"Total.*?AED\s+([\d,]+(?:\.\d+)?)", text, re.I | re.S)
    if m:
        data["total"] = float(m.group(1).replace(",", ""))

    # Bill-to customer
    m = re.search(r"Bill to[:\s]*\n\s*(.+?)(?:\n|,)", text, re.I)
    if m:
        data["customer_name"] = m.group(1).strip()

    # Address fields
    m = re.search(r"PO\s*Box\s+(\d+)", text, re.I)
    if m:
        data["bill_to_po_box"] = m.group(1)

    bill_addr_parts = []
    if "bill_to_po_box" in data:
        bill_addr_parts.append(f"PO Box {data['bill_to_po_box']}")
    city_m = re.search(r"(?:Dubai|Abu Dhabi|Sharjah|Ajman|RAK)", text, re.I)
    if city_m:
        bill_addr_parts.append(city_m.group(0))
    bill_addr_parts.append("UAE")
    data["bill_to_address"] = ", ".join(bill_addr_parts)

    # Vendor address
    addr_block = re.search(
        r"Qrestik Technologies\s+L\.?L\.?C\.?\s*\n(.+?)\nBill to",
        text, re.I | re.S
    )
    if addr_block:
        lines = [l.strip() for l in addr_block.group(1).splitlines() if l.strip()]
        data["ship_to_address"] = ", ".join(lines)

    # Banking (RAKBANK)
    data["bank_account_number"] = _first_match(r"Account Number\s*[–\-:]\s*([\d]+)", text)
    data["bank_iban"]            = _first_match(r"IBAN Number\s*[–\-:]\s*([A-Z0-9]+)", text)
    data["bank_branch"]          = _first_match(r"Branch Name\s*[–\-:]\s*(.+)", text)
    data["bank_swift"]           = _first_match(r"Swift Code\s*[–\-:]\s*([A-Z0-9]+)", text)
    data["bank_routing"]         = _first_match(r"Routing Code\s*[–\-:]\s*([0-9]+)", text)
    data["bank_address"]         = _first_match(r"Address\s*[–\-:]\s*(.+)", text)

    # Line items
    line_items: list[LineItemSchema] = []
    item_block = re.search(
        r"Description\s+Total Amount\s*\n(.+?)\nRemittance",
        text, re.I | re.S
    )
    if item_block:
        for row in item_block.group(1).splitlines():
            row = row.strip()
            if not row:
                continue
            lm = re.match(r"(.+?)\s+(AED|USD|EUR)\s+([\d,]+(?:\.\d+)?)", row, re.I)
            if lm:
                amt = float(lm.group(3).replace(",", ""))
                line_items.append(LineItemSchema(
                    description=lm.group(1).strip(),
                    qty=1, rate=amt, amount=amt
                ))
    # Fallback
    if not line_items:
        desc = _first_match(r"(Oracle Fusion[^\n]+)", text)
        if desc:
            amt_raw = _first_match(r"AED\s+([\d,]+(?:\.\d+)?)", text)
            amt = float(amt_raw.replace(",", "")) if amt_raw else 0.0
            line_items.append(LineItemSchema(description=desc, qty=1, rate=amt, amount=amt))
    data["line_items"] = line_items

    # Subtotal / tax
    if data.get("total"):
        data["subtotal"] = data["total"]
        data["tax"] = 0.0

    return data


# ---------------------------------------------------------------------------
# Infinitum-specific extraction
# ---------------------------------------------------------------------------

def _extract_infinitum(text: str) -> dict:
    data: dict = {}

    # Invoice number — "INVOICE#: 3611"
    m = re.search(r"INVOICE\s*#\s*[:\-]?\s*(\d+)", text, re.I)
    if m:
        data["invoice_number"] = m.group(1).strip()

    # Date — "DATE: 01st June 2026"
    m = re.search(r"DATE\s*[:\-]\s*(.+?)(?:\n|INVOICE)", text, re.I | re.S)
    if m:
        data["invoice_date"] = _parse_date(m.group(1).strip())

    data["currency"] = "USD"
    data["vendor_name"] = "Infinitum Global LLC"

    # Total
    m = re.search(r"Total\s+\$\s*([\d,]+(?:\.\d+)?)", text, re.I)
    if m:
        data["total"] = float(m.group(1).replace(",", ""))

    # Ship-to person name (first line after "Ship To")
    m = re.search(r"Ship To[:\s]*\n\s*(.+?)(?:\n)", text, re.I)
    if m:
        data["customer_name"] = m.group(1).strip()
    else:
        # Fallback to Bill-to
        m = re.search(r"Bill To[:\s]*\n\s*(.+?)(?:\n)", text, re.I)
        if m:
            data["customer_name"] = m.group(1).strip()

    bill_block = re.search(r"Bill To[:\s]*\n(.+?)(?:Ship To|Description)", text, re.I | re.S)
    if bill_block:
        lines = [l.strip() for l in bill_block.group(1).splitlines() if l.strip()]
        data["bill_to_address"] = "\n".join(lines[1:]) if len(lines) > 1 else None

    # Ship-to
    ship_block = re.search(r"Ship To[:\s]*\n(.+?)(?:Description|Po\.Number|$)", text, re.I | re.S)
    if ship_block:
        lines = [l.strip() for l in ship_block.group(1).splitlines() if l.strip()]
        ship_parts = []
        if len(lines) > 2:
            ship_parts = lines[2:]
        data["ship_to_address"] = "\n".join(ship_parts) if ship_parts else None

    # PO number — "#PO018558"
    m = re.search(r"#\s*(PO\d+)", text, re.I)
    if m:
        data["po_number"] = m.group(1)

    # Billing period
    m = re.search(
        r"(\d+\w*\s+\w+\s+\d{4})\s+to\s+(\d+\w*\s+\w+\s+\d{4})", text, re.I
    )
    period_start = m.group(1).strip() if m else None
    period_end   = m.group(2).strip() if m else None

    # Line items table
    line_items: list[LineItemSchema] = []
    item_block = re.search(
        r"Description\s+Po\.Number\s+Code/SKU\s+Period\s+Amount\s*\n(.+?)(?:Total|\Z)",
        text, re.I | re.S
    )
    if item_block:
        for row in item_block.group(1).splitlines():
            row = row.strip()
            if not row or row.lower().startswith("total"):
                continue
            lm = re.match(
                r"(.+?)\s+(#\s*\w+)\s+([\w\-]+)\s+"
                r"(\d+\w*\s+\w+\s+\d{4})\s+to\s+(\d+\w*\s+\w+\s+\d{4})\s+"
                r"\$\s*([\d,]+(?:\.\d+)?)",
                row
            )
            if lm:
                amt = float(lm.group(6).replace(",", ""))
                desc = f"{lm.group(1).strip()} | {lm.group(3).strip()}"
                if period_start:
                    desc += f" ({lm.group(4)} – {lm.group(5)})"
                line_items.append(LineItemSchema(description=desc, qty=1, rate=amt, amount=amt))

    # Fallback BRATS line
    if not line_items:
        bm = re.search(
            r"(BRATS\s+Project)\s+(#\s*\w+)\s+([\w\-]+)\s+"
            r"(.+?\d{4})\s+to\s+(.+?\d{4})\s+\$\s*([\d,]+(?:\.\d+)?)",
            text, re.I
        )
        if bm:
            amt = float(bm.group(6).replace(",", ""))
            line_items.append(LineItemSchema(
                description=f"{bm.group(1).strip()} | {bm.group(3).strip()} ({bm.group(4)} – {bm.group(5)})",
                qty=1, rate=amt, amount=amt
            ))
    data["line_items"] = line_items

    # Banking (JP Morgan)
    data["bank_name"]           = _first_match(r"Bank Name[:\s]+(.+)", text)
    data["bank_account_number"] = _first_match(r"Account[:\s]+([\d]+)", text)
    data["bank_routing"]        = _first_match(r"Routing\s*#[:\s]+([\d]+)", text)
    data["bank_fein"]           = _first_match(
        r"Federal Employee Identification Number\s*\(FEIN\)[:\s]+([\d\-]+)", text
    )
    data["bank_address"]        = _first_match(r"Address Associated w/\s*Account[:\s]+(.+)", text)
    data["bank_email"]          = _first_match(r"Email Associated w/\s*Account[:\s]+(\S+@\S+)", text)

    if data.get("total"):
        data["subtotal"] = data["total"]
        data["tax"] = 0.0

    return data


# ---------------------------------------------------------------------------
# Generic fallback extraction
# ---------------------------------------------------------------------------

def _extract_generic(text: str) -> dict:
    data: dict = {}

    inv_m = re.search(r"(?:INV[-\s]?|Invoice\s*#?\s*)(\d{4,}|[A-Z0-9-]{4,})", text, re.I)
    if inv_m:
        data["invoice_number"] = inv_m.group(0).strip().upper().replace(" ", "-")

    po_m = re.search(r"(?:PO|P\.O\.)\s*(?:#|No\.?)?\s*([A-Z0-9-]+)", text, re.I)
    if po_m:
        data["po_number"] = po_m.group(1)

    date_patterns = [
        r"(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})",
        r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{4})",
    ]
    for pat in date_patterns:
        m = re.search(r"Invoice\s*Date[:\s]*" + pat, text, re.I) or re.search(pat, text)
        if m:
            data["invoice_date"] = _parse_date(m.group(1))
            break

    money = re.findall(r"\$?\s*([\d,]+\.\d{2})", text)
    if money:
        amounts = [float(x.replace(",", "")) for x in money]
        data["total"] = max(amounts)
        if len(amounts) > 1:
            data["subtotal"] = sorted(amounts)[-2]
            data["tax"] = data["total"] - data["subtotal"]

    data["currency"] = "EUR" if "€" in text else "USD"

    # NER for customer name
    try:
        import spacy
        nlp = spacy.load("en_core_web_sm")
        doc = nlp(text[:100_000])
        orgs = [ent.text for ent in doc.ents if ent.label_ == "ORG"]
        if orgs:
            data["customer_name"] = orgs[0]
    except Exception:
        pass

    # Address blocks
    def _addr_block(header: str) -> Optional[str]:
        m = re.search(
            rf"{header}\s*[:\n]\s*(.+?)(?:\n\s*\n|Ship To|Bill To|Line Items|Subtotal|Total|$)",
            text, re.I | re.S
        )
        return m.group(1).strip()[:500] if m else None

    data["bill_to_address"] = _addr_block("Bill To")
    data["ship_to_address"] = _addr_block("Ship To")

    # Line items via table extraction
    line_items: list[LineItemSchema] = []
    try:
        import pdfplumber  # only available for PDFs
        # caller passes path via the text; not available here —
        # generic text-based line item fallback:
        pass
    except ImportError:
        pass
    for line in text.splitlines():
        lm = re.match(
            r"^(.+?)\s+(\d+(?:\.\d+)?)\s+\$?([\d,]+\.\d{2})\s+\$?([\d,]+\.\d{2})$",
            line.strip()
        )
        if lm:
            line_items.append(LineItemSchema(
                description=lm.group(1).strip(),
                qty=float(lm.group(2)),
                rate=float(lm.group(3).replace(",", "")),
                amount=float(lm.group(4).replace(",", "")),
            ))
    data["line_items"] = line_items
    return data


# ---------------------------------------------------------------------------
# Line item extraction from PDF tables (vendor-agnostic)
# ---------------------------------------------------------------------------

def _extract_line_items_pdf(path: str) -> list[LineItemSchema]:
    import pdfplumber
    largest: list[list] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            for table in (page.extract_tables() or []):
                if len(table) > len(largest):
                    largest = table
    items: list[LineItemSchema] = []
    if len(largest) > 1:
        for row in largest[1:]:
            if not row or not any(row):
                continue
            cells = [str(c or "").strip() for c in row]
            if len(cells) >= 4 and cells[0]:
                try:
                    qty = float(re.sub(r"[^\d.]", "", cells[1]) or 1)
                    rate = float(re.sub(r"[^\d.]", "", cells[2]) or 0)
                    amount = float(re.sub(r"[^\d.]", "", cells[3]) or qty * rate)
                    items.append(LineItemSchema(description=cells[0], qty=qty, rate=rate, amount=amount))
                except ValueError:
                    continue
    return items


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def parse_invoice(file_path: str) -> InvoiceParseSchema:
    """
    Auto-detect vendor from the uploaded PDF/DOCX and return a fully
    populated InvoiceParseSchema ready for front-end auto-fill.
    """
    path = Path(file_path)
    if not path.exists():
        return InvoiceParseSchema(missing_fields=REQUIRED_FIELDS + ["file"])

    ext = path.suffix.lower()
    if ext == ".pdf":
        text = _extract_text_pdf(file_path)
    elif ext in {".docx", ".doc"}:
        text = _extract_text_docx(file_path)
    else:
        return InvoiceParseSchema(missing_fields=REQUIRED_FIELDS + ["file_type"])

    vendor = _detect_vendor(text)
    logger.info("parse_invoice: detected vendor=%s for %s", vendor, path.name)

    if vendor == "qrestik":
        merged = _extract_qrestik(text)
        merged["vendor"] = "qrestik"
    elif vendor == "infinitum":
        merged = _extract_infinitum(text)
        merged["vendor"] = "infinitum"
    else:
        merged = _extract_generic(text)
        merged["vendor"] = "generic"
        # Try table extraction for generics
        if ext == ".pdf" and not merged.get("line_items"):
            merged["line_items"] = _extract_line_items_pdf(file_path)

    line_items = merged.pop("line_items", [])

    result = InvoiceParseSchema(
        raw_text_length=len(text),
        line_items=line_items,
        **{k: v for k, v in merged.items() if k not in ("line_items",) and v is not None},
    )
    result.missing_fields = [f for f in REQUIRED_FIELDS if getattr(result, f) in (None, "", [])]
    return result
